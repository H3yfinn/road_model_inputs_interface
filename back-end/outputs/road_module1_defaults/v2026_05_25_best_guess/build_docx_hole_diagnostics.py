from __future__ import annotations

from pathlib import Path
import re
import pandas as pd
import docx

ROOT = Path(r"c:\Users\Work\github\multinode_energy_balance\back-end\outputs\road_module1_defaults\v2026_05_25_best_guess")
DOCX_PATH = Path(r"C:\Users\Work\APERC\Outlook 10 - LEAP modelling_2026\Guides and notes\transport model process.docx")

WHITELIST_PATH = ROOT / "transport_model_process_branch_variable_whitelist.csv"
HOLES_PATH = ROOT / "road_module1_default_holes_expected_keys.csv"
GUIDANCE_OUT = ROOT / "transport_model_process_structured_guidance.csv"
HOLES_DOCX_OUT = ROOT / "road_module1_default_holes_docx_annotated.csv"

VARIABLES = [
    "Stock",
    "Sales Share",
    "Mileage",
    "Final On-Road Fuel Economy",
    "Vehicle Equivalent Weight",
    "Passenger Vehicle Saturation",
    "PHEV Electric Driving Share",
    "Reconciliation Bound Lower",
    "Reconciliation Bound Upper",
    "Reconciliation Weight",
    "Survival Rate",
    "Vintage Profile Share",
    "Sales",
    "Device Share",
    "Average Mileage",
    "Fuel Economy",
    "Final On-Road Mileage",
]

BRANCH_HINTS = [
    r"Demand\\Passenger road",
    r"Demand\\Freight road",
    "Passenger road",
    "Freight road",
    "LPVs",
    "LCVs",
    "Trucks",
    "Buses",
    "Motorcycles",
    "ICE",
    "HEV",
    "PHEV",
    "BEV",
    "FCEV",
    "Age",
]


def infer_expected_scope(rule_text: str) -> str:
    t = rule_text.lower()
    if "fuel branches" in t or "fuel branch" in t:
        return "fuel_branch"
    if "engine type" in t:
        return "engine_type"
    if "vehicle type" in t:
        return "vehicle_type"
    if "transport type" in t:
        return "transport_type"
    if "current accounts" in t:
        return "current_accounts_context"
    if "projection scenario" in t or "reference or target" in t:
        return "projection_context"
    if "lifecycle" in t or "vintage" in t or "age structure" in t:
        return "lifecycle_profile"
    return "general"


def extract_raw_texts(doc_path: Path) -> list[str]:
    d = docx.Document(str(doc_path))
    texts: list[str] = []

    for p in d.paragraphs:
        txt = (p.text or "").strip()
        if txt:
            texts.append(txt)

    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = (cell.text or "").strip()
                if txt:
                    texts.append(txt)

    return texts


def build_guidance_rows(texts: list[str]) -> pd.DataFrame:
    rows = []
    idx = 0
    for text in texts:
        vm = [v for v in VARIABLES if v.lower() in text.lower()]
        bm = [b for b in BRANCH_HINTS if b.lower() in text.lower()]
        path_like = "\\" in text and any(k in text for k in ["Demand", "road", "LPVs", "LCVs", "Trucks", "Buses", "Motorcycles"])
        if not (vm or bm or path_like):
            continue

        idx += 1
        confidence = "high" if (vm and (bm or path_like)) else ("medium" if (vm or bm) else "low")
        rows.append(
            {
                "rule_id": f"R{idx:04d}",
                "rule_text": text,
                "variables": " | ".join(vm),
                "branch_hints": " | ".join(bm),
                "contains_path_like_text": bool(path_like),
                "expected_scope": infer_expected_scope(text),
                "confidence": confidence,
            }
        )

    return pd.DataFrame(rows)


def annotate_holes(holes_df: pd.DataFrame, guidance_df: pd.DataFrame) -> pd.DataFrame:
    variable_to_rules: dict[str, list[str]] = {}
    variable_to_scopes: dict[str, set[str]] = {}

    for row in guidance_df.itertuples(index=False):
        vars_in_row = [v.strip() for v in str(row.variables).split("|") if v.strip()]
        for v in vars_in_row:
            variable_to_rules.setdefault(v, []).append(row.rule_id)
            variable_to_scopes.setdefault(v, set()).add(row.expected_scope)

    def classify(row: pd.Series) -> tuple[str, str, int, str]:
        variable = str(row.get("Variable", "")).strip()
        branch = str(row.get("Branch_Path", "")).strip()
        matched_rule_ids = variable_to_rules.get(variable, [])
        matched_scopes = sorted(variable_to_scopes.get(variable, set()))

        if re.search(r"Motorcycles\\(HEV|PHEV|FCEV)", branch):
            reason = "possible_invalid_vehicle_drive_combination"
        elif variable in {"Survival Rate", "Vintage Profile Share"}:
            reason = "likely_not_in_transport_export_lifecycle_data"
        elif variable in {"Reconciliation Bound Lower", "Reconciliation Bound Upper", "Reconciliation Weight", "Vehicle Equivalent Weight", "Passenger Vehicle Saturation"}:
            reason = "likely_model_factor_not_in_transport_export"
        elif variable in {"Stock", "Sales Share", "Final On-Road Fuel Economy", "Mileage"}:
            reason = "likely_branch_or_scope_mapping_mismatch"
        else:
            reason = "review_required"

        return (
            reason,
            " | ".join(matched_scopes),
            len(matched_rule_ids),
            ";".join(matched_rule_ids[:10]),
        )

    out = holes_df.copy()
    out[[
        "likely_docx_cause",
        "docx_expected_scopes",
        "docx_matched_rule_count",
        "docx_matched_rule_ids",
    ]] = out.apply(classify, axis=1, result_type="expand")

    return out


def main() -> None:
    texts = extract_raw_texts(DOCX_PATH)
    guidance_df = build_guidance_rows(texts)
    guidance_df.to_csv(GUIDANCE_OUT, index=False, encoding="utf-8-sig")

    holes_df = pd.read_csv(HOLES_PATH)
    annotated = annotate_holes(holes_df, guidance_df)
    annotated.to_csv(HOLES_DOCX_OUT, index=False, encoding="utf-8-sig")

    print(f"Wrote: {GUIDANCE_OUT}")
    print(f"Rows: {len(guidance_df)}")
    print(f"Wrote: {HOLES_DOCX_OUT}")
    print(f"Rows: {len(annotated)}")
    print("likely_docx_cause counts:")
    print(annotated["likely_docx_cause"].value_counts().to_string())


if __name__ == "__main__":
    main()
